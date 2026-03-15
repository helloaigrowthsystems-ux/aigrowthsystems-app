# exportador.py — AImpulsa v2
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

AZUL = colors.HexColor("#1E3A5F")
AZUL_SUAVE = colors.HexColor("#2E6DA4")
GRIS_LINEA = colors.HexColor("#D0D7E3")
GRIS_FONDO = colors.HexColor("#F4F6FA")
NEGRO = colors.HexColor("#1A1A2E")
BLANCO = colors.white

def _parsear_texto(texto):
    bloques = []
    for linea in texto.splitlines():
        s = linea.strip()
        if not s:
            bloques.append({"tipo": "espacio"})
        elif re.match(r'^(CLÁUSULA|ARTÍCULO|SECCIÓN|ANEXO|PARTE|CAPÍTULO)\s', s, re.I):
            bloques.append({"tipo": "h2", "texto": s})
        elif s.isupper() and 4 < len(s) < 80:
            bloques.append({"tipo": "h1", "texto": s})
        elif s.startswith(("- ", "• ", "* ")):
            bloques.append({"tipo": "bullet", "texto": s[2:]})
        elif re.match(r'^\d+[\.\)]\s', s):
            bloques.append({"tipo": "numerado", "texto": s})
        else:
            bloques.append({"tipo": "parrafo", "texto": s})
    return bloques

def exportar_pdf(contenido, tipo_doc="documento", metadata=None):
    buffer = BytesIO()
    metadata = metadata or {}
    doc = SimpleDocTemplate(buffer, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm, topMargin=2.5*cm, bottomMargin=2.5*cm)
    ests = getSampleStyleSheet()
    e_h1 = ParagraphStyle("h1", parent=ests["Normal"], fontSize=13, fontName="Helvetica-Bold",
        textColor=AZUL, spaceAfter=6, spaceBefore=14, leading=16)
    e_h2 = ParagraphStyle("h2", parent=ests["Normal"], fontSize=11, fontName="Helvetica-Bold",
        textColor=AZUL_SUAVE, spaceAfter=4, spaceBefore=10, leading=14)
    e_p = ParagraphStyle("p", parent=ests["Normal"], fontSize=10, fontName="Helvetica",
        textColor=NEGRO, spaceAfter=4, spaceBefore=2, leading=15, alignment=TA_JUSTIFY)
    e_b = ParagraphStyle("b", parent=ests["Normal"], fontSize=10, fontName="Helvetica",
        textColor=NEGRO, leftIndent=20, spaceAfter=3, leading=14)
    historia = []
    historia.append(Table([[Paragraph(
        f'<font color="white"><b>{tipo_doc.upper()}</b></font>',
        ParagraphStyle("cab", fontSize=14, fontName="Helvetica-Bold", textColor=BLANCO, alignment=TA_CENTER)
    )]], colWidths=[doc.width], style=TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), AZUL),
        ("TOPPADDING", (0,0), (-1,-1), 10), ("BOTTOMPADDING", (0,0), (-1,-1), 10),
    ])))
    historia.append(Spacer(1, 0.4*cm))
    filas = [(k, v) for k, v in [
        ("Empresa:", metadata.get("empresa","")), ("Cliente:", metadata.get("cliente","")),
        ("Fecha:", metadata.get("fecha","")), ("Número:", metadata.get("numero","")),
    ] if v]
    if filas:
        historia.append(Table(filas, colWidths=[3.5*cm, doc.width-3.5*cm], style=TableStyle([
            ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"), ("FONTNAME", (1,0), (1,-1), "Helvetica"),
            ("FONTSIZE", (0,0), (-1,-1), 9), ("TEXTCOLOR", (0,0), (0,-1), AZUL_SUAVE),
            ("TEXTCOLOR", (1,0), (1,-1), NEGRO), ("TOPPADDING", (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3), ("BACKGROUND", (0,0), (-1,-1), GRIS_FONDO),
            ("BOX", (0,0), (-1,-1), 0.5, GRIS_LINEA),
        ])))
        historia.append(Spacer(1, 0.5*cm))
    historia.append(HRFlowable(width="100%", thickness=1.5, color=AZUL))
    historia.append(Spacer(1, 0.4*cm))
    for bloque in _parsear_texto(contenido):
        t = bloque["tipo"]
        if t == "espacio":
            historia.append(Spacer(1, 0.25*cm))
        elif t == "h1":
            historia.append(Paragraph(bloque["texto"], e_h1))
            historia.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_LINEA))
        elif t == "h2":
            historia.append(Paragraph(bloque["texto"], e_h2))
        elif t in ("bullet", "numerado"):
            historia.append(Paragraph(f"&#8226;  {bloque['texto']}", e_b))
        else:
            historia.append(Paragraph(bloque["texto"], e_p))
    historia.append(Spacer(1, 1.5*cm))
    historia.append(HRFlowable(width="100%", thickness=0.5, color=GRIS_LINEA))
    historia.append(Spacer(1, 0.3*cm))
    w = doc.width/2 - 0.5*cm
    e_f = ParagraphStyle("f", fontSize=9, fontName="Helvetica", textColor=NEGRO, alignment=TA_CENTER)
    historia.append(Table([[
        Paragraph("Firma y sello<br/><br/><br/>___________________________<br/>Representante empresa", e_f),
        Paragraph("Firma y sello<br/><br/><br/>___________________________<br/>Cliente / Receptor", e_f),
    ]], colWidths=[w, w], style=TableStyle([
        ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6)
    ])))
    doc.build(historia)
    buffer.seek(0)
    return buffer.read()

def _set_cell_bg(celda, hex_color):
    tc_pr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def exportar_docx(contenido, tipo_doc="documento", metadata=None):
    metadata = metadata or {}
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    _set_cell_bg(c, "1E3A5F")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tipo_doc.upper())
    r.font.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(8)
    doc.add_paragraph()
    campos = [(k,v) for k,v in [
        ("Empresa:", metadata.get("empresa","")), ("Cliente:", metadata.get("cliente","")),
        ("Fecha:", metadata.get("fecha","")), ("Número:", metadata.get("numero","")),
    ] if v]
    if campos:
        tm = doc.add_table(rows=len(campos), cols=2)
        tm.style = "Table Grid"
        for i, (k, v) in enumerate(campos):
            ck, cv = tm.cell(i,0), tm.cell(i,1)
            _set_cell_bg(ck, "F4F6FA"); _set_cell_bg(cv, "F4F6FA")
            rk = ck.paragraphs[0].add_run(k); rk.bold = True
            rk.font.color.rgb = RGBColor(0x2E,0x6D,0xA4); rk.font.size = Pt(9)
            rv = cv.paragraphs[0].add_run(v); rv.font.size = Pt(9)
        doc.add_paragraph()
    for bloque in _parsear_texto(contenido):
        t2 = bloque["tipo"]
        if t2 == "espacio":
            doc.add_paragraph()
        elif t2 == "h1":
            p2 = doc.add_heading(bloque["texto"], level=1)
            if p2.runs: p2.runs[0].font.color.rgb = RGBColor(0x1E,0x3A,0x5F); p2.runs[0].font.size = Pt(13)
        elif t2 == "h2":
            p2 = doc.add_heading(bloque["texto"], level=2)
            if p2.runs: p2.runs[0].font.color.rgb = RGBColor(0x2E,0x6D,0xA4); p2.runs[0].font.size = Pt(11)
        elif t2 in ("bullet", "numerado"):
            doc.add_paragraph(bloque["texto"], style="List Bullet")
        else:
            p2 = doc.add_paragraph(bloque["texto"])
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p2.runs: run.font.size = Pt(10)
    doc.add_paragraph()
    tf = doc.add_table(rows=1, cols=2)
    for i, txt in enumerate(["Representante empresa", "Cliente / Receptor"]):
        cf = tf.cell(0, i)
        cf.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cf.paragraphs[0].add_run(f"Firma y sello\n\n\n___________________________\n{txt}").font.size = Pt(9)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()